import os
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

# Ensure a directory exists for temporary files 📂
os.makedirs("temp_downloads", exist_ok=True)

def create_docx(text, filename="digitized_note"):
    """Generates a standard Word document."""
    path = f"temp_downloads/{filename}_{datetime.now().strftime('%H%M%S')}.docx"
    doc = Document()
    doc.add_heading('Digitized Study Notes', 0)
    
    # Adding text in standard paragraphs
    paragraphs = text.split('\n')
    for p_text in paragraphs:
        if p_text.strip():
            doc.add_paragraph(p_text.strip())
            
    doc.save(path)
    return path

def create_pdf(text, filename="digitized_note"):
    """Generates a professional PDF with 1-inch margins and Times-Roman font."""
    path = f"temp_downloads/{filename}_{datetime.now().strftime('%H%M%S')}.pdf"
    
    # 1. Setup the Document Template (72 points = 1 inch) 📐
    doc = SimpleDocTemplate(
        path,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    # 2. Define Custom Styles 🎨
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontName='Times-Bold',
        fontSize=18,
        alignment=1,      # Centered
        spaceAfter=24     # Gap after title
    )
    
    essay_style = ParagraphStyle(
        'EssayStyle',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=12,
        leading=16,       # Professional line spacing
        spaceAfter=12,    # Gap between paragraphs
        alignment=0       # Left aligned
    )
    
    # 3. Build the "Story" (Content) ✍️
    story = []
    story.append(Paragraph("Digitized Study Notes", title_style))
    
    # Split text into paragraphs based on LLM output
    paragraphs = text.split('\n')
    for p_text in paragraphs:
        if p_text.strip():
            # Paragraph handles automatic word-wrapping within the margins
            story.append(Paragraph(p_text.strip(), essay_style))
    
    # 4. Create the File
    doc.build(story)
    return path